import hashlib
import re
from pathlib import Path
from typing import List, Dict
import pdfplumber
from docx import Document
from .config import RAGS_DIR, SUPPORTED_LANGUAGES

def is_arabic(text: str) -> bool:
    """Check if text contains Arabic characters."""
    arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+')
    return bool(arabic_pattern.search(text))

def detect_language(text: str) -> str:
    """Detect if text is primarily Arabic or English."""
    return "ar" if is_arabic(text) else "en"

def clean_text(text: str, language: str) -> str:
    """Clean and normalize text based on language."""
    if not text:
        return ""
    text = re.sub(r'\s+', ' ', text).strip()
    if language == "ar":
        # Remove Arabic diacritics
        text = re.sub(r'[ًٌٍَُِّْ]', '', text)
    return text

def load_multiple_pdfs() -> List[Dict]: #It loads multiple pdfs or docx files.
    """
    Extract content from all PDF and DOCX files in the RAGS_DIR directory.
    Returns a list of message dicts with source tracking.
    """
    print(f"[DEBUG] Loading documents from: {RAGS_DIR}")
    if not RAGS_DIR.exists():
        print(f"[ERROR] Directory not found: {RAGS_DIR}")
        return []

    messages: List[Dict] = []
    files = list(RAGS_DIR.glob("*.pdf")) + list(RAGS_DIR.glob("*.docx"))
    if not files:
        print(f"[ERROR] No PDF or DOCX files in {RAGS_DIR}")
        return []

    print(f"[INFO] Found {len(files)} documents (.pdf/.docx)")

    for path in files:
        name = path.name
        suffix = path.suffix.lower()
        print(f"[INFO] Processing {name}")

        try:
            if suffix == ".pdf":
                # PDF extraction
                with pdfplumber.open(path) as pdf:
                    for page_idx, page in enumerate(pdf.pages):
                        raw = page.extract_text() or ""
                        lang = detect_language(raw)
                        paras = re.split(r'\n\s*\n', raw)
                        for i, para in enumerate(paras):
                            para = para.strip()
                            if not para:
                                continue
                            cleaned = clean_text(para, lang)
                            cid = hashlib.md5(f"{name}_p{page_idx}_para{i}:{cleaned}".encode()).hexdigest()
                            hdr = re.match(r'^([\w\s\u0600-\u06FF]+)[:|-]', para)
                            section = hdr.group(1).strip() if hdr else "General"
                            is_table = bool(re.search(r'\n\s+\w+.*\n\s+\w+', para))
                            messages.append({
                                "id": cid,
                                "source_file": name,
                                "page": page_idx+1,
                                "sender": "SFDA_Doc",
                                "section": section,
                                "is_table": is_table,
                                "language": lang,
                                "text": cleaned
                            })
                        # PDF tables
                        for t_idx, tbl in enumerate(page.extract_tables()):
                            if not tbl:
                                continue
                            text_tbl = "\n".join([" | ".join(cell or "" for cell in row) for row in tbl])
                            lang_tbl = detect_language(text_tbl)
                            cleaned_tbl = clean_text(text_tbl, lang_tbl)
                            tid = hashlib.md5(f"{name}_p{page_idx}_tbl{t_idx}:{cleaned_tbl}".encode()).hexdigest()
                            messages.append({
                                "id": tid,
                                "source_file": name,
                                "page": page_idx+1,
                                "sender": "SFDA_Table",
                                "section": "Table Data",
                                "is_table": True,
                                "language": lang_tbl,
                                "text": cleaned_tbl
                            })

            elif suffix == ".docx":
                # DOCX extraction
                doc = Document(path)
                for para_idx, para in enumerate(doc.paragraphs):
                    raw = para.text or ""
                    if not raw.strip():
                        continue
                    lang = detect_language(raw)
                    cleaned = clean_text(raw, lang)
                    cid = hashlib.md5(f"{name}_para{para_idx}:{cleaned}".encode()).hexdigest()
                    hdr = re.match(r'^([\w\s\u0600-\u06FF]+)[:|-]', raw)
                    section = hdr.group(1).strip() if hdr else "General"
                    messages.append({
                        "id": cid,
                        "source_file": name,
                        "page": None,
                        "sender": "SFDA_Docx",
                        "section": section,
                        "is_table": False,
                        "language": lang,
                        "text": cleaned
                    })
                # DOCX tables
                for tbl_idx, table in enumerate(doc.tables):
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(" | ".join(cells))
                    text_tbl = "\n".join(rows)
                    if not text_tbl.strip():
                        continue
                    lang_tbl = detect_language(text_tbl)
                    cleaned_tbl = clean_text(text_tbl, lang_tbl)
                    tid = hashlib.md5(f"{name}_tbl{tbl_idx}:{cleaned_tbl}".encode()).hexdigest()
                    messages.append({
                        "id": tid,
                        "source_file": name,
                        "page": None,
                        "sender": "SFDA_Table",
                                "section": "Table Data",
                                "is_table": True,
                                "language": lang_tbl,
                                "text": cleaned_tbl
                    })
            else:
                print(f"[WARN] Unsupported file type: {name}")
        except Exception as err:
            print(f"[ERROR] Failed {name}: {err}")

    print(f"[INFO] Extracted {len(messages)} chunks")
    ar_count = sum(1 for m in messages if m.get("language") == "ar")
    en_count = sum(1 for m in messages if m.get("language") == "en")
    print(f"[INFO] Distribution: Arabic={ar_count}, English={en_count}")
    return messages
